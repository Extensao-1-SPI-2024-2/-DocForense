import json
from django.http import JsonResponse
from django.contrib.auth import authenticate, login, logout
from .services import GalileuAPIService
from django.http import HttpResponse
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from .models import Modelo, ModeloVestigio, Variavel, Chamado, RespostaChamado
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib import messages
from .template_processor import TemplateProcessor
from .contexto_procedimento_pericial import ContextoProcedimentoPericial
from datetime import datetime
from django.http import HttpResponse
from docx import Document
from bs4 import BeautifulSoup


@login_required(login_url='/login')
def procedimento_pericial(request, id):
    if request.method == 'GET':
        service = GalileuAPIService()
        response_data = service.get_procedimento_pericial(id)

        if 'error' in response_data:
            return JsonResponse({'status': 'error', 'message': response_data['error']}, status=500)

        return JsonResponse({'status': 'success', 'data': response_data})
    else:
        return JsonResponse({'error': 'Method not allowed'}, status=405)


@login_required(login_url='/login')
def gerar_arquivo_word_laudo(request):
    if request.method == 'POST':
        html_content = request.POST.get('laudo', '')

        if not html_content:
            return HttpResponse("Nenhum HTML foi enviado.", status=400)

        doc = Document()

        soup = BeautifulSoup(html_content, 'html.parser')

        for tag in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li']):
            if tag.name.startswith('h'):
                doc.add_heading(tag.get_text(), level=int(tag.name[1]))
            elif tag.name == 'p':
                doc.add_paragraph(tag.get_text())
            elif tag.name in ['ul', 'ol']:
                for li in tag.find_all('li'):
                    doc.add_paragraph(f'• {li.get_text()}', style='List Bullet')

        from io import BytesIO
        file_buffer = BytesIO()
        doc.save(file_buffer)
        file_buffer.seek(0)

        response = HttpResponse(
            file_buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="laudo.docx"'
        return response

    return HttpResponse("Método não permitido.", status=405)


def logininternal(request):
    if request.method == 'GET':
        if 'logged_in' in request.session and request.session['logged_in']:
            return redirect('/')

    if request.method == 'POST':
        # Criar o usuário admin, se não existir
        user, created = User.objects.get_or_create(
            username='admindoc1',
            defaults={
                'email': 'admin@docforense.com.br',
                'first_name': 'Master',
                'last_name': 'Administrador',
            }
        )

        if created:
            user.set_password('admindoc1')
            user.save()

        # Criar o usuário comum, se não existir
        user_comum, criado_comum = User.objects.get_or_create(
            username='comumuser',
            defaults={
                'email': 'comumuser@exemplo.com',
                'first_name': 'Comum',
                'last_name': 'User',
            }
        )

        if criado_comum:
            user_comum.set_password('comumuser')
            user_comum.save()

        # Autenticar o usuário que fez o login
        username = request.POST.get('username')
        password = request.POST.get('password')

        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            request.session['logged_in'] = True
            request.session['login'] = username
            return redirect('/')
        else:
            messages.error(request, "Usuário ou senha incorretos.")

    return render(request, 'login.html')


@login_required(login_url='/login')
def index(request):
    if request.method == 'GET':
        if not request.session.get('logged_in', False):
            messages.error(request, "É necessário realizar o login.")
            return redirect('/login')

    return render(request, 'index.html')


@login_required(login_url='/login')
def logoutinternal(request):
    request.session.pop('logged_in', None)
    request.session.pop('login', None)
    logout(request)
    return redirect('/login')


@login_required(login_url='/login')
def create_modelo(request):
    if request.method == 'POST':

        name = request.POST.get('name')
        value = request.POST.get('value')
        type = request.POST.get('type')

        if not all([name, value, type]):
            messages.error(request, "Todos os campos são obrigatórios.")
            return redirect('/modelo/laudo/criar')

        try:
            Modelo.objects.create(
                name=name,
                value=value,
                type=type,
                user_inclusion=request.user
            )

            messages.success(request, "Modelo criado com sucesso!")
            return redirect('/modelo/laudo/listagem')
        except User.DoesNotExist:
            messages.error(request, "Usuário não encontrado")
            return redirect('/modelo/laudo/criar')
        except Exception as e:
            messages.error(request, f"Erro: {str(e)}")
            return redirect('/modelo/laudo/criar')

    variaveis = Variavel.objects.filter(tipo=0).values('variavel', 'descricao')
    vestigios = ModeloVestigio.objects.values('id', 'name', 'type_vestigio')

    vestigios_organizados = {}
    for vestigio in vestigios:
        type_vestigio = vestigio['type_vestigio']
        if type_vestigio not in vestigios_organizados:
            vestigios_organizados[type_vestigio] = []
        vestigios_organizados[type_vestigio].append(vestigio)

    return render(request, 'modelo.html', {
        'variaveis': json.dumps(list(variaveis)),
        'type_choices': ModeloVestigio.TYPE_CHOICES,
        'vestigios': json.dumps(vestigios_organizados),
        'type_vestigio_choices': json.dumps(dict(ModeloVestigio.TYPE_VESTIGIO_CHOICES))
    })


@login_required(login_url='/login')
def create_modelo_vestigio(request, type):
    tipo_descricao = next((desc for value, desc in Variavel.TIPO_CHOICES if value == type), None)

    if not tipo_descricao:
        messages.error(request, "Esse tipo de vestígio não foi encontrado.")
        return redirect('/modelo/vestigio/listagem')

    if request.method == 'POST':

        name = request.POST.get('name')
        value = request.POST.get('value')
        type_view = request.POST.get('type')

        if not all([name, value, type_view]):
            messages.error(request, "Todos os campos são obrigatórios.")
            return redirect('/modelo/vestigio/criar')

        try:
            ModeloVestigio.objects.create(
                name=name,
                value=value,
                type=type_view,
                type_vestigio=type,
                user_inclusion=request.user
            )

            messages.success(request, "Modelo de vestígio criado com sucesso!")
            return redirect('/modelo/vestigio/listagem')
        except User.DoesNotExist:
            messages.error(request, "Usuário não encontrado")
            return redirect('/modelo/vestigio/criar')
        except Exception as e:
            messages.error(request, f"Erro: {str(e)}")
            return redirect('/modelo/vestigio/criar')

    variaveis = Variavel.objects.filter(tipo=type).values('variavel', 'descricao')

    return render(request, 'modelo_vestigio.html', {
        'variaveis': json.dumps(list(variaveis)),
        'tipo': type,
        'tipo_descricao': tipo_descricao,
        'type_choices': ModeloVestigio.TYPE_CHOICES
    })


@login_required(login_url='/login')
def listagem(request):
    if request.method == 'GET':
        if not request.session.get('logged_in', False):
            messages.error(request, "É necessário realizar o login.")
            return redirect('/login')

    modelos = Modelo.objects.all()

    return render(request, 'listagem.html', {'modelos': modelos})


@login_required(login_url='/login')
def listagemVestigio(request):
    if request.method == 'GET':
        if not request.session.get('logged_in', False):
            messages.error(request, "É necessário realizar o login.")
            return redirect('/login')

    modelos = ModeloVestigio.objects.all()

    return render(request, 'listagem_vestigio.html', {
        'modelos': modelos
    })


@login_required(login_url='/login')
def editar_modelo(request, modelo_id):
    modelo = get_object_or_404(Modelo, id=modelo_id)

    if request.method == 'POST':
        name = request.POST.get('name')
        value = request.POST.get('value')
        type_view = request.POST.get('type')

        if not all([name, value, type_view]):
            messages.error(request, "Todos os campos são obrigatórios.")
            return redirect(f'/modelo/laudo/editar/{modelo_id}')

        try:
            modelo.name = name
            modelo.value = value
            modelo.type = type_view
            modelo.save()

            messages.success(request, "Modelo atualizado com sucesso!")
            return redirect('/modelo/laudo/listagem')
        except Exception as e:
            messages.error(request, f"Erro ao atualizar o modelo: {str(e)}")
            return redirect(f'/modelo/laudo/editar/{modelo_id}')

    variaveis = Variavel.objects.filter(tipo=0).values('variavel', 'descricao')
    vestigios = ModeloVestigio.objects.values('id', 'name', 'type_vestigio')

    vestigios_organizados = {}
    for vestigio in vestigios:
        type_vestigio = vestigio['type_vestigio']
        if type_vestigio not in vestigios_organizados:
            vestigios_organizados[type_vestigio] = []
        vestigios_organizados[type_vestigio].append(vestigio)

    return render(request, 'modelo_edicao.html', {
        'variaveis': json.dumps(list(variaveis)),
        'modelo': modelo,
        'readonly': False,
        'type_choices': ModeloVestigio.TYPE_CHOICES,
        'vestigios': json.dumps(vestigios_organizados),
        'type_vestigio_choices': json.dumps(dict(ModeloVestigio.TYPE_VESTIGIO_CHOICES))
    })


@login_required(login_url='/login')
def editar_modelo_vestigio(request, modelo_id):
    modelo = get_object_or_404(ModeloVestigio, id=modelo_id)

    if request.method == 'POST':
        name = request.POST.get('name')
        value = request.POST.get('value')
        type_view = request.POST.get('type')

        if not all([name, value, type_view]):
            messages.error(request, "Todos os campos são obrigatórios.")
            return redirect(f'/modelo/vestigio/editar/{modelo_id}')

        try:
            modelo.name = name
            modelo.value = value
            modelo.type = type_view
            modelo.save()

            messages.success(request, "Modelo de Vestígio atualizado com sucesso!")
            return redirect('/modelo/vestigio/listagem')
        except Exception as e:
            messages.error(request, f"Erro ao atualizar o modelo de vestígio: {str(e)}")
            return redirect(f'/modelo/vestigio/editar/{modelo_id}')

    variaveis = Variavel.objects.filter(tipo=modelo.type_vestigio).values('variavel', 'descricao')

    return render(request, 'modelo_vestigio_edicao.html', {
        'modelo': modelo,
        'readonly': False,
        'variaveis': json.dumps(list(variaveis)),
        'tipo': modelo.type_vestigio,
        'type_choices': ModeloVestigio.TYPE_CHOICES,
        'tipo_descricao': modelo.get_type_vestigio_display()
    })


@login_required(login_url='/login')
def visualizar_modelo(request, modelo_id):
    modelo = get_object_or_404(Modelo, id=modelo_id)
    vestigios = ModeloVestigio.objects.values('id', 'name', 'type_vestigio')

    vestigios_organizados = {}
    for vestigio in vestigios:
        type_vestigio = vestigio['type_vestigio']
        if type_vestigio not in vestigios_organizados:
            vestigios_organizados[type_vestigio] = []
        vestigios_organizados[type_vestigio].append(vestigio)

    return render(request, 'modelo_edicao.html', {
        'modelo': modelo,
        'readonly': True,
        'type_choices': ModeloVestigio.TYPE_CHOICES,
        'vestigios': json.dumps(vestigios_organizados),
        'type_vestigio_choices': json.dumps(dict(ModeloVestigio.TYPE_VESTIGIO_CHOICES))
    })


@login_required(login_url='/login')
def visualizar_modelo_vestigio(request, modelo_id):
    modelo = get_object_or_404(ModeloVestigio, id=modelo_id)
    return render(request, 'modelo_vestigio_edicao.html', {
        'modelo': modelo,
        'readonly': True,
        'type_choices': ModeloVestigio.TYPE_CHOICES,
        'tipo_descricao': modelo.get_type_vestigio_display()
    })


@login_required(login_url='/login')
def deletar_modelo(request, modelo_id):
    modelo = get_object_or_404(Modelo, id=modelo_id)

    try:
        modelo.delete()
        messages.success(request, "Modelo deletado com sucesso!")
    except Exception as e:
        messages.error(request, f"Erro ao deletar o modelo: {str(e)}")

    return redirect('/modelo/laudo/listagem')


@login_required(login_url='/login')
def deletar_modelo_vestigio(request, modelo_id):
    modelo = get_object_or_404(ModeloVestigio, id=modelo_id)

    try:
        modelo.delete()
        messages.success(request, "Modelo de Vestígio deletado com sucesso!")
    except Exception as e:
        messages.error(request, f"Erro ao deletar o modelo de vestígio: {str(e)}")

    return redirect('/modelo/vestigio/listagem')


@login_required(login_url='/login')
def gerar_laudo(request):
    modelos = Modelo.objects.all()
    return render(request, 'gerar_laudo.html', {'modelos': modelos})


def converter_timestamp_para_data_brasileira(timestamp_ms):
    if timestamp_ms is None:
        return None

    timestamp = timestamp_ms / 1000

    dt = datetime.fromtimestamp(timestamp)

    return dt.strftime('%d/%m/%Y %H:%M:%S')


@login_required(login_url='/login')
def gerar_modelo_formatado(request, galileu_id, modelo_id):
    if request.method == 'GET':

        service = GalileuAPIService()
        response_data = service.get_procedimento_pericial(galileu_id)

        if 'error' in response_data:
            return JsonResponse({'status': 'error', 'message': response_data['error']}, status=500)

        contexto = ContextoProcedimentoPericial(response_data).gerar_contexto()

        modelo = get_object_or_404(Modelo, id=modelo_id)

        processor = TemplateProcessor(contexto)
        text = processor.processar_vestigios(modelo.value, response_data)
        text = processor.substituir_variaveis(text)

        return JsonResponse({'status': 'success', 'data': {
            "text": text
        }})

    return JsonResponse({'error': 'Method not allowed'}, status=405)


@login_required(login_url='/login')
def listagem_chamado(request):
    if request.user.is_superuser:
        chamados = Chamado.objects.all()
    else:
        chamados = Chamado.objects.filter(user_inclusion=request.user)
    
    return render(request, 'listagem_chamado.html', {
        'is_super_user': request.user.is_superuser,
        'chamados': chamados
    })

@login_required(login_url='/login')
def gerar_chamado(request):
    if request.user.is_superuser:
        messages.error(request, "Usuário administrador não autorizado a gerar chamados.")
        return redirect('/chamado/listagem')
    
    if request.method == 'POST':
        assunto = request.POST.get('assunto')
        urgencia = request.POST.get('urgencia')
        descricao = request.POST.get('descricao')

        if not all([assunto, urgencia, descricao]):
            messages.error(request, "Todos os campos são obrigatórios.")
            return redirect(f'/chamado/criar')

        try:
            Chamado.objects.create(
                assunto=assunto,
                descricao=descricao,
                responsavel=2,
                status=1,
                urgencia=urgencia,
                user_inclusion=request.user
            )
            messages.success(request, "Chamado criado com sucesso!")
            return redirect('/chamado/listagem')
        except Exception as e:
            messages.error(request, f"Erro ao criar o chamado: {str(e)}")
            return redirect('/chamado/criar')

    return render(request, 'chamado.html', {
        'urgencia_choices': Chamado.URGENCIA_CHOICES
    })

@login_required(login_url='/login')
def analisar_chamado(request, chamado_id):
    chamado = get_object_or_404(Chamado, id=chamado_id)

    if not request.user.is_superuser and request.user != chamado.user_inclusion:
        messages.error(request, "Acesso não autorizado.")
        return redirect('/chamado/listagem')
    
    historico = RespostaChamado.objects.filter(chamado=chamado).order_by('-date_inclusion')

    return render(request, 'chamado_analise.html', {
        'is_super_user': request.user.is_superuser,
        'chamado': chamado,
        'urgencia_choices': Chamado.URGENCIA_CHOICES,
        'historico': historico
    })

@login_required(login_url='/login')
def responder_chamado(request, chamado_id):
    chamado = get_object_or_404(Chamado, id=chamado_id)

    if chamado.responsavel == 2 and not request.user.is_superuser:
        messages.error(request, "Este chamado somente pode ser respondido por administradores.")
        return redirect('/chamado/listagem')
    
    if chamado.responsavel == 1 and request.user != chamado.user_inclusion:
        messages.error(request, "Este chamado somente pode ser respondido pelo solicitante.")
        return redirect('/chamado/listagem')
    
    if chamado.status == 3:
        messages.error(request, "Este chamado se encontra encerrado, por favor cadastre um novo.")
        return redirect('/chamado/listagem')
    
    if request.method == 'POST':
        try:
            resposta = request.POST.get('resposta')

            if not all([resposta]):
                messages.error(request, "Todos os campos são obrigatórios.")
                return redirect(f'/chamado/' + chamado_id)
        
            RespostaChamado.objects.create(
                chamado=chamado,
                mensagem=resposta,
                autor=request.user
            )

            chamado.responsavel = 1 if chamado.responsavel == 2 else 2
            chamado.status = 2
            chamado.save()
            
            messages.success(request, "Chamado respondido com sucesso.")
            return redirect('/chamado/listagem')
        except Exception as e:
            messages.error(request, f"Erro ao responder o chamado: {str(e)}")
            return redirect('/chamado/' + chamado_id)
    else:
        messages.error(request, "Requisição não autorizada.")
        return redirect('/chamado/listagem')